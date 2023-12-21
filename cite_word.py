import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

#assume at this point the excel is correct
#and has at least the columns "date" in yyyy-mm-dd, "url", "title", "author", no weird html text anywhere, etc
#sorting groups:
#group by publisher, and within publisher, sort by date. If same date then technically by author last name but idrc

#bullet bold: publisher
#indent bullet: author name. "Title." <i>publisher</i>, date, url.
#*for date, format %d %b %Y, add . after month if abbreviated (i.e. not May, June, or July)

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class ParagraphExt:
	#From https://github.com/python-openxml/python-docx/pull/582#issuecomment-1717139576
    p = None

    def __init__(self, p): #p: paragraph obj
        self.p = p

    def restart_numbering(self):
        """
        Restarting the numbering of paragraph
        """

        # Getting the abstract number of paragraph
        abstract_num_id = self.p.part.document.part.numbering_part.element.num_having_numId(
            self.p.style.element.get_or_add_pPr().get_or_add_numPr().numId.val).abstractNumId.val

        # Add abstract number to numbering part and reset
        num = self.p.part.numbering_part.element.add_num(abstract_num_id)
        num.add_lvlOverride(ilvl=0).add_startOverride(1)

        # Get or add elements to paragraph
        p_pr = self.p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = num_pr.get_or_add_ilvl()
        ilvl.val = int("0")
        num_id = num_pr.get_or_add_numId()
        num_id.val = int(num.numId)

def format_date(date):
	#in: yyyy-mm-dd
	#out: dd mon., yyyy
	# _, m, _ = date.split("-")
	# month = int(m)
	mydate = datetime.strptime(date, "%Y-%m-%d")
	#mydate = date.to_pydatetime()
	if 5 <= mydate.month and mydate.month <= 7:
		ret = datetime.strftime(mydate, "%d %b %Y")
	else:
		ret = datetime.strftime(mydate, "%d %b. %Y")
	return(ret)

def add_citation_authorless(doc, author, title, publisher, date, url, level=2):
	#to write
	if(level == 1):
		mystyle = 'ListNumber'
	else:
		mystyle = 'ListNumber'+str(level)
	out = doc.add_paragraph()
	out.style = mystyle
	run = out.add_run()
	if(len(author) > 0):
		run.add_text(author+'. "'+title+'."')
	else:
		run.add_text('"'+title+'."')
	
	publisher_run = out.add_run(publisher)
	publisher_run.italic = True
	publisher_run.add_text(', ')

	run2 = out.add_run(format_date(date))
	run2.add_text(', '+url)

def add_citation(doc, reset_numbering, author, title, publisher, date, url):
	out = doc.add_paragraph()
	out.style = 'List Number 2'
	run = out.add_run()
	if(len(author.strip()) > 0):
		run.add_text(author+'. "'+title+'."')
	else:
		run.add_text('"'+title+'."')
	
	publisher_run = out.add_run(publisher)
	publisher_run.italic = True
	publisher_run.add_text(', ')

	run2 = out.add_run(format_date(date))
	run2.add_text(', '+url)

	if(reset_numbering):
		p_ext = ParagraphExt(out)
		p_ext.restart_numbering()

def main(infile, outfile, titletxt):
	raw = pd.read_excel(infile)
	raw = raw.fillna('')

	doc = Document()

	title = doc.add_paragraph()
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	title_run = title.add_run()
	title_font = title_run.font
	title_font.size = Pt(14)
	title_run.bold = True
	title_run.add_text(titletxt)

	blank = doc.add_paragraph()
	blank.alignment = WD_ALIGN_PARAGRAPH.LEFT

	publishers = sorted(set(raw['publisher']))

	for p in publishers:
		subset = raw[raw['publisher'] == p]
		pub_header = doc.add_paragraph()
		pub_header.style = "List Bullet"
		pub_run = pub_header.add_run(p)
		pub_run.add_text(' ('+str(subset.shape[0])+')')
		pub_run.bold = True
		new_list = True
		for i in range(subset.shape[0]):
			add_citation(doc, new_list, *subset.iloc[i, [subset.columns.get_loc(col) for col in ['authors', 'title', 'publisher', 'date', 'url']]])
			new_list = False


	doc.save(outfile)
	print("wrote to", outfile)
	return

main('product/BP-2022-10-01-2023-09-30.xlsx', 'product/BP-2022-10-01-2023-09-30.docx', 'Beyond Parallel Media Coverage 10/2022 - 9/2023')